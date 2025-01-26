import os
import time
import pytesseract
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import WebDriverException
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from PIL import Image
import json
from docx import Document
import pyfiglet
from fpdf import FPDF
import pandas as pd
import csv
import sys
from colorama import Fore, Back, Style


def create_results_folder(folder_name):
    """Creates the folder for saving results if it doesn't exist."""
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
        print(Fore.GREEN + f"\n[INFO] Folder '{folder_name}' created.")
    else:
        print(Fore.YELLOW + f"\n[INFO] Folder '{folder_name}' already exists.")


def display_banner():
    """Displays a banner showcasing the web scraper tool's name and functionality."""
    terminal_width = os.get_terminal_size().columns
    font = "slant"  # Choose a bold and impactful font for a hacking look

    # Generate the banner with pyfiglet
    banner = pyfiglet.figlet_format(Fore.BLUE + "WEBSITE DATA SCRAPER", font=font, width=terminal_width)
    banner_lines = banner.splitlines()  # Split the banner into lines

    # Center-align each line of the banner based on terminal width
    centered_banner = "\n".join(line.center(terminal_width) for line in banner_lines)

    # Display the banner and additional details
    print(Fore.LIGHTGREEN_EX + centered_banner)
    print(Fore.YELLOW + "Script Written by Dev Elijah".center(terminal_width))
    print(Fore.CYAN + "Scrape text, images, and videos from websites.".center(terminal_width))
    print(Fore.RED + "=" * terminal_width)  # Full-width separator


def spinner_animation():
    """Display a spinning animation while scraping."""
    spin_chars = ['|', '/', '-', '\\']
    for i in range(10):  # Limit the loop to show a spinning effect a few times
        for char in spin_chars:
            sys.stdout.write(f'\r{Fore.GREEN}\nScraping {char}')
            sys.stdout.flush()
            time.sleep(0.1)
    sys.stdout.write(f'\r{Fore.GREEN}\nScraping Complete!       \n')


def detect_website_info(soup):
    """Extract basic information about the website using OG meta tags or the title."""
    meta_description = soup.find("meta", {"name": "description"}) or soup.find("meta", {"property": "og:description"})
    site_info = {
        "title": soup.title.text.strip() if soup.title else "No Title Found",
        "description": meta_description.get("content").strip() if meta_description else "No Description Found"
    }
    return site_info


def extract_text_from_image(image_path):
    """Extract text from an image using Tesseract OCR."""
    try:
        text = pytesseract.image_to_string(Image.open(image_path))
        return text
    except Exception as e:
        print(Fore.RED + f"\nError extracting text from image: {e}")
        return ""


def download_image(driver, image_url, save_path):
    """Download and save an image from the website."""
    try:
        driver.get(image_url)
        screenshot_path = save_path if save_path.endswith('.png') else f"{save_path}.png"
        driver.save_screenshot(screenshot_path)
        return screenshot_path
    except Exception as e:
        print(Fore.RED + f"\nError downloading image {image_url}: {e}")
        return None


def print_progress_bar(iteration, total, prefix='', suffix='', length=40, fill='█', print_end="\r"):
    """Print a real-time progress bar."""
    percent = ("{0:.1f}").format(100 * (iteration / float(total)))
    filled_length = int(length * iteration // total)
    bar = fill * filled_length + '-' * (length - filled_length)
    sys.stdout.write(f'\r{prefix} |{bar}| {percent}% {suffix}')
    sys.stdout.flush()


page_times = []  # List to store the time taken for the last few pages

def scrape_page(driver, url, visited_urls, data_collected, keyword, scrape_images, scrape_videos, scrape_text_from_images, scrape_all_pages, page_number, total_pages):
    """Scrape a single page and extract relevant information."""

    # Start time for this page
    page_start_time = time.time()
    
    if url in visited_urls:
        return

    visited_urls.add(url)
    try:
        driver.get(url)
        time.sleep(3)
    except WebDriverException as e:
        print(Fore.RED + f"\nError accessing {url}: {e}")
        return

    soup = BeautifulSoup(driver.page_source, 'html.parser')
    site_info = detect_website_info(soup)
    print(Fore.YELLOW + f"\nWebsite Title: {site_info['title']}")
    print(Fore.YELLOW + f"\nWebsite Description: {site_info['description']}\n")

    # Extract text and save only those containing the keyword
    if keyword:
        page_text = soup.get_text()
        if keyword.lower() in page_text.lower():
            data_collected['text'].append({"url": url, "content": page_text})
            print(Fore.GREEN + f"\nScraped Text (Keyword Found) from {url}")
    else:
        data_collected['text'].append({"url": url, "content": soup.get_text()})
        print(Fore.GREEN + f"\nScraped Text from {url}")

    # Extract images
    if scrape_images:
        for img in soup.find_all('img'):
            src = img.get('src')
            if src:
                image_url = urljoin(url, src)
                image_path = download_image(driver, image_url, f"{len(visited_urls)}_image.png")
                if scrape_text_from_images and image_path:
                    text_from_image = extract_text_from_image(image_path)
                    data_collected['images'].append({'image_url': image_url, 'extracted_text': text_from_image})
                    print(Fore.YELLOW + f"\nScraped Image and Extracted Text from {url}")
                else:
                    data_collected['images'].append({'image_url': image_url})
                    print(Fore.YELLOW + f"\nScraped Image from {url}")

    # Extract videos
    if scrape_videos:
        for video in soup.find_all('video'):
            src = video.get('src')
            if src:
                video_url = urljoin(url, src)
                data_collected['videos'].append(video_url)
                print(Fore.CYAN + f"\nScraped Video from {url}")

    # Calculate the time taken for this page
    page_end_time = time.time()
    page_time_taken = page_end_time - page_start_time

    # Add the time taken for this page to the list of page times
    page_times.append(page_time_taken)

    # If there are more than 5 times stored, remove the oldest one
    if len(page_times) > 5:
        page_times.pop(0)

    # Calculate the average time taken for the last few pages
    average_page_time = sum(page_times) / len(page_times)

    # Estimate the remaining time
    pages_left = total_pages - page_number
    estimated_remaining_time = pages_left * average_page_time
    estimated_remaining_minutes = estimated_remaining_time / 60

    print(Fore.GREEN + f"\nTime taken for page {page_number}: {page_time_taken:.2f} seconds.")
    print(Fore.CYAN + f"\nEstimated time remaining: {estimated_remaining_minutes:.2f} minutes.")

    # Update progress bar after scraping a page
    print_progress_bar(page_number, total_pages, prefix='\nScraping Pages', suffix='\nComplete', length=50)

    # Follow links (if scraping all pages)
    if scrape_all_pages:
        for link in soup.find_all('a', href=True):
            link_url = urljoin(url, link['href'])
            if urlparse(link_url).netloc == urlparse(url).netloc:
                scrape_page(driver, link_url, visited_urls, data_collected, keyword, scrape_images, scrape_videos, scrape_text_from_images, scrape_all_pages, page_number+1, total_pages)


def save_data(data, folder_name, file_format, keyword=None):
    """Save scraped data in the specified format."""
    os.makedirs(folder_name, exist_ok=True)  # This ensures the folder exists, but doesn't recreate it.
    file_path = os.path.join(folder_name, f"scraped_data.{file_format}")

    if file_format == "json":
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
    elif file_format == "txt":
        with open(file_path, "w", encoding="utf-8") as f:
            for section, items in data.items():
                f.write(f"--- {section.upper()} ---\n")
                for item in items:
                    if section == 'text' and keyword:
                        # Only save content containing the keyword
                        if keyword.lower() in item['content'].lower():
                            f.write(f"URL: {item.get('url', 'N/A')}\nContent: {item.get('content', '')}\n\n")
                    else:
                        f.write(f"URL: {item.get('url', 'N/A')}\nContent: {item.get('content', '')}\n\n")
    elif file_format == "docx":
        doc = Document()
        for section, items in data.items():
            doc.add_heading(section.upper(), level=1)
            for item in items:
                if section == 'text' and keyword:
                    # Only save content containing the keyword
                    if keyword.lower() in item['content'].lower():
                        doc.add_paragraph(f"URL: {item.get('url', 'N/A')}")
                        doc.add_paragraph(item.get('content', ''))
                else:
                    doc.add_paragraph(f"URL: {item.get('url', 'N/A')}")
                    doc.add_paragraph(item.get('content', ''))
        doc.save(file_path)
    elif file_format == "csv":
        with open(file_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Section", "URL", "Content"])
            for section, items in data.items():
                for item in items:
                    if section == 'text' and keyword:
                        # Only save content containing the keyword
                        if keyword.lower() in item['content'].lower():
                            writer.writerow([section, item.get('url', 'N/A'), item.get('content', '')])
                    else:
                        writer.writerow([section, item.get('url', 'N/A'), item.get('content', '')])
    elif file_format == "pdf":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for section, items in data.items():
            pdf.cell(200, 10, f"--- {section.upper()} ---", ln=True, align="C")
            for item in items:
                if section == 'text' and keyword:
                    # Only save content containing the keyword
                    if keyword.lower() in item['content'].lower():
                        pdf.multi_cell(0, 10, f"URL: {item.get('url', 'N/A')}")
                        pdf.multi_cell(0, 10, item.get('content', ''))
                else:
                    pdf.multi_cell(0, 10, f"URL: {item.get('url', 'N/A')}")
                    pdf.multi_cell(0, 10, item.get('content', ''))
        pdf.output(file_path)
    elif file_format == "xlsx":
        data_for_excel = []
        for section, items in data.items():
            for item in items:
                data_for_excel.append([section, item.get('url', 'N/A'), item.get('content', '')])
        df = pd.DataFrame(data_for_excel, columns=["Section", "URL", "Content"])
        df.to_excel(file_path, index=False)
    else:
        print(Fore.RED + f"\nUnsupported format: {file_format}")
        return

    print(Fore.LIGHTGREEN_EX + f"Data saved in folder: {folder_name}, Format: {file_format.upper()}")


def scrape_website():
    """Main function to scrape an entire website."""
    display_banner()

    while True:
        url = input(Fore.CYAN + "Enter the URL of the website to scrape: ").strip()
        # Add 'https://' if the URL doesn't already have it
        if not url.lower().startswith(('http://', 'https://')):
            url = 'https://' + url
        try:
            result = urlparse(url)
            if all([result.scheme, result.netloc]):
                break
            else:
                raise ValueError
        except ValueError:
            print(Fore.RED + "\nInvalid URL. Please enter a valid URL.")

    keyword = input(Fore.YELLOW + "\nEnter the keyword to search for (leave empty to scrape all content): ").strip()
    folder_name = input(Fore.GREEN + "\nEnter the folder name to save the results: ").strip()
    file_format = input(Fore.BLUE + "\nEnter the file format to save data (json, txt, docx, csv, pdf, xlsx): ").strip().lower()
    scrape_images = input(Fore.MAGENTA + "\nDo you want to scrape images? (yes/no): ").strip().lower() == 'yes'
    scrape_videos = input(Fore.MAGENTA + "\nDo you want to scrape videos? (yes/no): ").strip().lower() == 'yes'
    scrape_text_from_images = input(Fore.MAGENTA + "\nDo you want to scrape text from images? (yes/no): ").strip().lower() == 'yes'
    
    # Ask about scraping specific page or all pages
    scrape_specific_page = input(Fore.MAGENTA + "\nDo you want to scrape a specific page? (yes/no): ").strip().lower()
    if scrape_specific_page == 'yes':
        specific_url = input(Fore.CYAN + "\nEnter the URL of the specific page to scrape: ").strip()
        scrape_all_pages = False
        url = specific_url
    else:
        scrape_all_pages = input(Fore.MAGENTA + "\nDo you want to scrape all pages of the site? (yes/no): ").strip().lower() == 'yes'
    
    open_browser = input(Fore.MAGENTA + "\nDo you want to open the browser while scraping? (yes/no): ").strip().lower() == 'yes'

    # Start measuring time
    start_time = time.time()

    create_results_folder(folder_name)

    # Configure the browser driver options
    chrome_options = Options()
    if not open_browser:
        chrome_options.add_argument("--headless")

    driver = webdriver.Chrome(options=chrome_options)

    data_collected = {'text': [], 'images': [], 'videos': []}
    visited_urls = set()

    # Start the scraping process
    spinner_animation()

    try:
        # Determine how many pages to scrape (this part can be adjusted to suit specific needs)
        total_pages = 10  # Example placeholder for total number of pages

        # Calculate elapsed time and print it
        end_time = time.time()
        elapsed_time = end_time - start_time
        print(Fore.CYAN + f"\nScraping completed in {elapsed_time:.2f} seconds.")

        scrape_page(driver, url, visited_urls, data_collected, keyword, scrape_images, scrape_videos, scrape_text_from_images, scrape_all_pages, 1, total_pages)
        save_data(data_collected, folder_name, file_format, keyword)
    except Exception as e:
        print(Fore.RED + f"\nError: {e}")
    finally:
        driver.quit()


if __name__ == "__main__":
    scrape_website()
